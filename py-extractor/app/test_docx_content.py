from __future__ import annotations

import io
import unittest

from docx import Document

from app import main


class DocxContentTest(unittest.TestCase):
    def test_paragraph_and_table_blocks_keep_body_order(self):
        document = Document()
        document.add_paragraph("before-table")
        first = document.add_table(rows=2, cols=2)
        first.cell(0, 0).text = "H|1"
        first.cell(0, 1).text = "H2"
        first.cell(1, 0).text = "row"
        first.cell(1, 1).text = "x"
        document.add_paragraph("")
        document.add_paragraph("between-tables")
        second = document.add_table(rows=2, cols=1)
        second.cell(0, 0).text = "Second"
        second.cell(1, 0).text = "value"
        document.add_paragraph("after-table")

        buffer = io.BytesIO()
        document.save(buffer)
        extracted = main.extract_docx(buffer.getvalue(), "fixture.docx")

        first_table = extracted.text.index(r"H\|1")
        second_table = extracted.text.index("| Second |")
        self.assertLess(extracted.text.index("before-table"), first_table)
        self.assertLess(first_table, extracted.text.index("between-tables"))
        self.assertLess(extracted.text.index("between-tables"), second_table)
        self.assertLess(second_table, extracted.text.index("after-table"))
        self.assertEqual(extracted.paragraph_count, 3)
        self.assertEqual(extracted.table_count, 2)


if __name__ == "__main__":
    unittest.main()
